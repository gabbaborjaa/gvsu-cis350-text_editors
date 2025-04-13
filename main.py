import eel
import os
from fpdf import FPDF
from docx import Document
from htmldocx import HtmlToDocx
import flet as ft

class FileManager:
    def __init__(self):
        eel.init("src")


    @eel.expose
    @staticmethod
    def new_file():
        eel.start('index.html')

    @eel.expose
    def save_file(self, filename, content, filetype, override=False):
        # Add the correct file extension if not already present
        if not filename.endswith(f".{filetype}"):
            filename = f"{filename}.{filetype}"

        if filetype not in ["txt", "md", "pdf", "docx"]:
            return f"Unsupported file type: {filetype}"

        if os.path.exists(filename) and not override:
            return f"File {filename} already exists. Do you want to override it?"

        try:
            if filetype == 'pdf':
                self._save_as_pdf(filename, content)
            elif filetype == 'md':
                self._save_as_text(filename, content)
            elif filetype == 'txt':
                self._save_as_text(filename, content)
            elif filetype == 'docx':
                self._save_as_docx(filename, content)
            else:
                return f"Unsupported file type: {filetype}"
            return f"File {filename} saved successfully."
        except Exception as e:
            return f"An error occurred while saving the file: {str(e)}"

    def _save_as_pdf(self, filename, content):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, content)
        pdf.output(filename)

    def _save_as_text(self, filename, content):
        with open(filename, 'w') as file:
            file.write(content)

    def _save_as_docx(self, filename, content):
        doc = Document()
        converter = HtmlToDocx()
        converter.add_html_to_document(content, doc)
        doc.save(filename)

    @eel.expose
    def load_file(self, filename):
        """
        Loads the content of a file and returns it to the frontend.
        Supports .txt, .md, and .docx file types.
        """
        if not os.path.exists(filename):
            return f"File {filename} does not exist."

        try:
            file_extension = filename.split('.')[-1].lower()

            if file_extension == "txt" or file_extension == "md":
                # Read plain text or markdown files
                with open(filename, 'r', encoding='utf-8') as file:
                    content = file.read()
                return content

            elif file_extension == "docx":
                # Read .docx files using python-docx
                doc = Document(filename)
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return content

            else:
                return f"Unsupported file type: {file_extension}"

        except Exception as e:
            return f"An error occurred while loading the file: {str(e)}"

    @eel.expose
    def undo(self):
        pass

    @eel.expose
    def redo(self):
        pass

    def start_app(self):
        eel.start('index.html')


if __name__ == "__main__":
    file_manager = FileManager()
    file_manager.start_app()
